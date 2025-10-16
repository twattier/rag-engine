"""Microsoft Word parser using python-docx."""
from __future__ import annotations

from pathlib import Path
from docx import Document

from app.parsers.base_parser import BaseParser
from app.models import ContentItem


class DOCXParser(BaseParser):
    """Parser for Microsoft Word (.docx) files."""

    async def parse(self, file_path: Path) -> list[ContentItem]:
        """Parse DOCX file and return content."""
        doc = Document(file_path)
        content_list: list[ContentItem] = []

        # Extract paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        if full_text:
            content_list.append(
                ContentItem(
                    type="text",
                    text="\n".join(full_text),
                    page_idx=0,
                    structure="docx_paragraphs",
                )
            )

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                rows.append(row_data)

            if rows:
                content_list.append(
                    ContentItem(
                        type="table",
                        rows=rows,
                        page_idx=0,
                        structure=f"docx_table_{table_idx}",
                    )
                )

        return content_list
